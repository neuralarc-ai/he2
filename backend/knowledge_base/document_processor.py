"""
Document Processor Service

This module handles the extraction, processing, and storage of documents
from various file formats into the knowledge base system.
"""

import asyncio
import io
import json
import logging
from typing import Dict, Any, Optional, List, Tuple
from pathlib import Path
import tempfile
import os

# Document processing libraries
try:
    import PyPDF2
    import pandas as pd
    from docx import Document
    import markdown
    from bs4 import BeautifulSoup
    EXTRACTORS_AVAILABLE = True
except ImportError:
    EXTRACTORS_AVAILABLE = False
    logging.warning("Some document extractors not available. Install required packages for full functionality.")

from utils.logger import logger

class DocumentProcessor:
    """
    Service for processing and extracting text from various document formats
    """
    
    def __init__(self):
        self.supported_formats = {
            'application/pdf': self._extract_pdf_text,
            'text/csv': self._extract_csv_text,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx_text,
            'text/plain': self._extract_text_text,
            'text/markdown': self._extract_markdown_text,
            'application/json': self._extract_json_text,
            'text/html': self._extract_html_text,
            'application/vnd.ms-excel': self._extract_excel_text,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._extract_excel_text,
        }
        
        if not EXTRACTORS_AVAILABLE:
            logger.warning("Document extractors not fully available. Some formats may not work.")
    
    async def process_document(
        self, 
        file_content: bytes, 
        mime_type: str, 
        filename: str,
        chunk_size: int = 1000,
        overlap: int = 200
    ) -> Dict[str, Any]:
        """
        Process a document and extract text content with chunking
        
        Args:
            file_content: Raw file bytes
            mime_type: MIME type of the file
            filename: Original filename
            chunk_size: Maximum size of each text chunk
            overlap: Overlap between chunks for context continuity
            
        Returns:
            Dictionary containing processed text and metadata
        """
        try:
            logger.info(f"Processing document: {filename} ({mime_type})")
            
            # Extract text based on file type
            if mime_type in self.supported_formats:
                extractor = self.supported_formats[mime_type]
                raw_text = await extractor(file_content, filename)
            else:
                logger.warning(f"Unsupported file type: {mime_type}")
                raw_text = self._extract_fallback_text(file_content, filename)
            
            if not raw_text or not raw_text.strip():
                logger.warning(f"No text extracted from {filename}")
                return {
                    "success": False,
                    "error": "No text content could be extracted",
                    "chunks": [],
                    "total_chunks": 0,
                    "metadata": self._extract_metadata(file_content, mime_type, filename)
                }
            
            # Clean and normalize text
            cleaned_text = self._clean_text(raw_text)
            
            # Chunk the text for better LLM consumption
            chunks = self._chunk_text(cleaned_text, chunk_size, overlap)
            
            # Extract metadata
            metadata = self._extract_metadata(file_content, mime_type, filename)
            
            logger.info(f"Successfully processed {filename}: {len(chunks)} chunks created")
            
            return {
                "success": True,
                "raw_text": raw_text,
                "cleaned_text": cleaned_text,
                "chunks": chunks,
                "total_chunks": len(chunks),
                "metadata": metadata,
                "total_tokens": len(cleaned_text) // 4,  # Rough token estimation
                "processing_info": {
                    "chunk_size": chunk_size,
                    "overlap": overlap,
                    "original_size": len(raw_text),
                    "cleaned_size": len(cleaned_text)
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {str(e)}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "chunks": [],
                "total_chunks": 0,
                "metadata": {}
            }
    
    async def _extract_pdf_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from PDF files"""
        if not EXTRACTORS_AVAILABLE:
            return self._extract_fallback_text(file_content, filename)
        
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Could not extract text from page {page_num + 1}: {e}")
                    continue
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting PDF text from {filename}: {e}")
            return self._extract_fallback_text(file_content, filename)
    
    async def _extract_csv_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from CSV files"""
        if not EXTRACTORS_AVAILABLE:
            return self._extract_fallback_text(file_content, filename)
        
        try:
            # Try to decode as UTF-8 first
            try:
                csv_text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                csv_text = file_content.decode('latin-1')
            
            # Parse CSV and convert to readable format
            df = pd.read_csv(io.StringIO(csv_text))
            
            # Convert to markdown table format for better LLM consumption
            markdown_table = df.to_markdown(index=False)
            
            # Add column descriptions
            columns_info = f"Columns: {', '.join(df.columns.tolist())}\n"
            rows_info = f"Total rows: {len(df)}\n\n"
            
            return f"{columns_info}{rows_info}{markdown_table}"
            
        except Exception as e:
            logger.error(f"Error extracting CSV text from {filename}: {e}")
            return self._extract_fallback_text(file_content, filename)
    
    async def _extract_docx_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from DOCX files"""
        if not EXTRACTORS_AVAILABLE:
            return self._extract_fallback_text(file_content, filename)
        
        try:
            doc = Document(io.BytesIO(file_content))
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract table data
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    text_parts.append("Table:\n" + "\n".join(table_text))
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text from {filename}: {e}")
            return self._extract_fallback_text(file_content, filename)
    
    async def _extract_text_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from plain text files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            return file_content.decode('utf-8', errors='ignore')
            
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            return self._extract_fallback_text(file_content, filename)
    
    async def _extract_markdown_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from markdown files"""
        if not EXTRACTORS_AVAILABLE:
            return self._extract_text_text(file_content, filename)
        
        try:
            # First extract as text
            raw_text = await self._extract_text_text(file_content, filename)
            
            # Convert markdown to HTML then extract text
            html = markdown.markdown(raw_text)
            soup = BeautifulSoup(html, 'html.parser')
            
            # Extract text while preserving some structure
            text_parts = []
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'code', 'pre']):
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    text_parts.append(f"\n{element.get_text().strip()}\n")
                elif element.name == 'code':
                    text_parts.append(f"`{element.get_text().strip()}`")
                elif element.name == 'pre':
                    text_parts.append(f"\n```\n{element.get_text().strip()}\n```\n")
                else:
                    text_parts.append(element.get_text().strip())
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting markdown text from {filename}: {e}")
            return await self._extract_text_text(file_content, filename)
    
    async def _extract_json_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from JSON files"""
        try:
            # Parse JSON
            json_data = json.loads(file_content.decode('utf-8'))
            
            # Convert to readable format
            if isinstance(json_data, dict):
                return self._format_json_dict(json_data)
            elif isinstance(json_data, list):
                return self._format_json_list(json_data)
            else:
                return str(json_data)
                
        except Exception as e:
            logger.error(f"Error extracting JSON text from {filename}: {e}")
            return self._extract_fallback_text(file_content, filename)
    
    async def _extract_html_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from HTML files"""
        if not EXTRACTORS_AVAILABLE:
            return self._extract_text_text(file_content, filename)
        
        try:
            html_text = file_content.decode('utf-8', errors='ignore')
            soup = BeautifulSoup(html_text, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting HTML text from {filename}: {e}")
            return await self._extract_text_text(file_content, filename)
    
    async def _extract_excel_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from Excel files"""
        if not EXTRACTORS_AVAILABLE:
            return self._extract_fallback_text(file_content, filename)
        
        try:
            # Read Excel file
            df = pd.read_excel(io.BytesIO(file_content), sheet_name=None)
            
            text_parts = []
            for sheet_name, sheet_df in df.items():
                if not sheet_df.empty:
                    text_parts.append(f"--- Sheet: {sheet_name} ---")
                    text_parts.append(f"Columns: {', '.join(sheet_df.columns.tolist())}")
                    text_parts.append(f"Rows: {len(sheet_df)}")
                    text_parts.append("Data:")
                    text_parts.append(sheet_df.to_markdown(index=False))
                    text_parts.append("")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting Excel text from {filename}: {e}")
            return self._extract_fallback_text(file_content, filename)
    
    def _extract_fallback_text(self, file_content: bytes, filename: str) -> str:
        """Fallback text extraction for unsupported formats"""
        try:
            # Try to decode as text
            return file_content.decode('utf-8', errors='ignore')
        except:
            # If all else fails, return a placeholder
            return f"[Binary file content from {filename} - text extraction not available]"
    
    def _format_json_dict(self, data: Dict[str, Any], indent: int = 0) -> str:
        """Format JSON dictionary for readable text"""
        lines = []
        for key, value in data.items():
            if isinstance(value, dict):
                lines.append(f"{'  ' * indent}{key}:")
                lines.append(self._format_json_dict(value, indent + 1))
            elif isinstance(value, list):
                lines.append(f"{'  ' * indent}{key}:")
                lines.append(self._format_json_list(value, indent + 1))
            else:
                lines.append(f"{'  ' * indent}{key}: {value}")
        return "\n".join(lines)
    
    def _format_json_list(self, data: List[Any], indent: int = 0) -> str:
        """Format JSON list for readable text"""
        lines = []
        for i, item in enumerate(data):
            if isinstance(item, dict):
                lines.append(f"{'  ' * indent}[{i}]:")
                lines.append(self._format_json_dict(item, indent + 1))
            elif isinstance(item, list):
                lines.append(f"{'  ' * indent}[{i}]:")
                lines.append(self._format_json_list(item, indent + 1))
            else:
                lines.append(f"{'  ' * indent}[{i}]: {item}")
        return "\n".join(lines)
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Remove excessive spaces
            cleaned_line = ' '.join(line.split())
            if cleaned_line.strip():
                cleaned_lines.append(cleaned_line)
        
        # Join lines with proper spacing
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Remove excessive newlines
        while '\n\n\n' in cleaned_text:
            cleaned_text = cleaned_text.replace('\n\n\n', '\n\n')
        
        return cleaned_text.strip()
    
    def _chunk_text(self, text: str, chunk_size: int, overlap: int) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks for better LLM consumption"""
        if not text:
            return []
        
        chunks = []
        words = text.split()
        current_chunk = []
        current_size = 0
        
        for i, word in enumerate(words):
            word_size = len(word) + 1  # +1 for space
            
            if current_size + word_size > chunk_size and current_chunk:
                # Create chunk
                chunk_text = ' '.join(current_chunk)
                chunks.append({
                    "text": chunk_text,
                    "size": len(chunk_text),
                    "word_count": len(current_chunk),
                    "start_word": i - len(current_chunk),
                    "end_word": i - 1
                })
                
                # Start new chunk with overlap
                overlap_words = current_chunk[-overlap:] if overlap > 0 else []
                current_chunk = overlap_words
                current_size = sum(len(w) + 1 for w in overlap_words)
            
            current_chunk.append(word)
            current_size += word_size
        
        # Add final chunk
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append({
                "text": chunk_text,
                "size": len(chunk_text),
                "word_count": len(current_chunk),
                "start_word": len(words) - len(current_chunk),
                "end_word": len(words) - 1
            })
        
        return chunks
    
    def _extract_metadata(self, file_content: bytes, mime_type: str, filename: str) -> Dict[str, Any]:
        """Extract metadata from the document"""
        return {
            "filename": filename,
            "mime_type": mime_type,
            "file_size_bytes": len(file_content),
            "file_size_mb": round(len(file_content) / (1024 * 1024), 2),
            "extension": Path(filename).suffix.lower(),
            "has_content": len(file_content) > 0,
            "processing_timestamp": None  # Will be set by caller
        }
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported MIME types"""
        return list(self.supported_formats.keys())
    
    def is_format_supported(self, mime_type: str) -> bool:
        """Check if a MIME type is supported"""
        return mime_type in self.supported_formats


# Async wrapper for the document processor
async def process_document_async(
    file_content: bytes,
    mime_type: str,
    filename: str,
    chunk_size: int = 1000,
    overlap: int = 200
) -> Dict[str, Any]:
    """Async wrapper for document processing"""
    processor = DocumentProcessor()
    return await processor.process_document(file_content, mime_type, filename, chunk_size, overlap)





